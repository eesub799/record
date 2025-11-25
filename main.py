from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import qrcode
from PIL import Image
import os
from typing import List
import io
import shutil

app = FastAPI(title="Lab Record Generator API")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class Experiment(BaseModel):
    title: str
    date: str = ""
    github: str

class RecordData(BaseModel):
    course_title: str
    student_name: str
    register_number: str
    experiments: List[Experiment]


def create_qr_code(url: str, size: int = 150):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=2,
    )
    qr.add_data(url)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")
    img = img.resize((size, size), Image.Resampling.LANCZOS)

    img_bytes = io.BytesIO()
    img.save(img_bytes, format="PNG")
    img_bytes.seek(0)
    return img_bytes


@app.get("/")
async def root():
    return {
        "message": "Lab Record Generator API",
        "status": "running",
        "version": "1.0",
        "logo_uploaded": os.path.exists("college_logo.png")
    }


@app.post("/generate-docx")
async def generate_docx(data: RecordData):
    try:
        doc = Document()

        # Basic layout
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(data.course_title)
        run.bold = True
        run.font.size = Pt(14)

        # Table
        table = doc.add_table(rows=len(data.experiments) + 1, cols=3)
        headers = ["Exp", "Title", "GitHub"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for idx, exp in enumerate(data.experiments):
            row = table.rows[idx + 1].cells
            row[0].text = str(idx + 1).zfill(2)
            row[1].text = exp.title
            row[2].text = exp.github

        filename = f"{data.register_number}_Lab_Record.docx"
        doc.save(filename)

        return FileResponse(
            filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health():
    return {"status": "healthy"}


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
